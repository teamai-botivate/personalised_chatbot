"""
Botivate HR Support - RAG (Document Search) Service
Indexes company policy documents and text rules into a vector store.
Answers are generated ONLY from company-specific data.
"""

import os
from typing import List, Optional
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain_community.vectorstores import Chroma
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_core.messages import HumanMessage, SystemMessage
from app.config import settings


def _get_collection_name(company_id: str) -> str:
    """Each company gets its own isolated vector collection."""
    return f"company_{company_id}"


def _get_embeddings() -> OpenAIEmbeddings:
    return OpenAIEmbeddings(api_key=settings.openai_api_key)


def _get_vectorstore(company_id: str) -> Chroma:
    return Chroma(
        collection_name=_get_collection_name(company_id),
        embedding_function=_get_embeddings(),
        persist_directory=settings.chroma_persist_dir,
    )


async def index_text_policy(company_id: str, title: str, content: str) -> None:
    """Index a text-based policy into the company's vector store."""
    if not settings.openai_api_key or settings.openai_api_key == "your-openai-api-key-here":
        print(f"⚠️ [MOCK RAG] Skipping actual AI embedding for policy '{title}' because OpenAI key is missing or dummy.")
        return

    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    chunks = splitter.split_text(content)

    documents = [
        Document(
            page_content=chunk,
            metadata={"company_id": company_id, "source": title, "type": "text_policy"},
        )
        for chunk in chunks
    ]

    vectorstore = _get_vectorstore(company_id)
    vectorstore.add_documents(documents)


async def index_document_file(company_id: str, title: str, file_path: str) -> None:
    """Extract text from a PDF/DOC and index it into the company's vector store."""
    text = ""

    if file_path.lower().endswith(".pdf"):
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        for page in reader.pages:
            extracted = page.extract_text()
            if extracted:
                text += extracted + "\n"

    elif file_path.lower().endswith(".docx"):
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"

    else:
        # Fallback: try reading as plain text
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
        except Exception:
            text = ""

    if text.strip():
        await index_text_policy(company_id, title, text)


async def search_policies(company_id: str, query: str, top_k: int = 5) -> List[Document]:
    """Search the company's policy vector store for relevant documents."""
    vectorstore = _get_vectorstore(company_id)
    results = vectorstore.similarity_search(query, k=top_k)
    return results


async def answer_from_policies(company_id: str, question: str) -> str:
    """
    Answer a question using ONLY the company's indexed policies.
    Never uses generic HR knowledge.
    """
    relevant_docs = await search_policies(company_id, question)

    if not relevant_docs:
        return "I could not find any relevant information in your company's policies. Please contact your HR for clarification."

    context = "\n\n---\n\n".join([doc.page_content for doc in relevant_docs])

    if settings.fast_llm_api_key and settings.fast_llm_model:
        kwargs = {
            "model": settings.fast_llm_model,
            "api_key": settings.fast_llm_api_key,
            "temperature": 0.1,
        }
        if settings.fast_llm_base_url:
            kwargs["base_url"] = settings.fast_llm_base_url
        llm = ChatOpenAI(**kwargs)
    else:
        llm = ChatOpenAI(
            model=settings.openai_model,
            api_key=settings.openai_api_key,
            temperature=0.1,
        )

    system_msg = SystemMessage(content="""You are an HR assistant that answers ONLY based on the provided company policy documents.

Rules:
- Answer STRICTLY from the context provided below.
- If the answer is NOT in the context, say "I could not find this information in your company's policies."
- NEVER guess, assume, or use generic HR knowledge.
- Be professional, clear, and helpful.
- Quote the policy title/source when relevant.
""")

    user_msg = HumanMessage(content=f"""Context from company policies:
{context}

Employee Question: {question}
""")

    response = await llm.ainvoke([system_msg, user_msg])
    return response.content.strip()
